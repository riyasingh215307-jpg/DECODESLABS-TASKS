from docx import Document
from docx.shared import Inches

import os# Create Word Document
document = Document()

# Title
document.add_heading("Decode Labs Project 3", level=1)

# Subtitle
document.add_heading(
    "Customer Segmentation using PCA and K-Means Clustering",
    level=2
)

# Student Details
document.add_paragraph("Name: Riya Singh")
document.add_paragraph("Internship: Decode Labs - Data Science Internship")
# ====================================================
# Introduction
# ====================================================

document.add_heading("1. Introduction", level=1)

document.add_paragraph(
    "Customer segmentation is the process of dividing customers into "
    "different groups based on their purchasing behaviour and characteristics. "
    "This project uses Principal Component Analysis (PCA) and K-Means Clustering "
    "to identify meaningful customer segments for business decision making."
)

# ====================================================
# Objective
# ====================================================

document.add_heading("2. Objective", level=1)

document.add_paragraph(
    "- Perform data preprocessing and cleaning.\n"
    "- Perform Exploratory Data Analysis (EDA).\n"
    "- Apply Feature Scaling.\n"
    "- Apply Principal Component Analysis (PCA).\n"
    "- Determine the optimal number of clusters using the Elbow Method.\n"
    "- Evaluate clustering using the Silhouette Score.\n"
    "- Perform Customer Segmentation using K-Means Clustering.\n"
    "- Generate business insights."
)

# ====================================================
# Dataset Description
# ====================================================

document.add_heading("3. Dataset Description", level=1)

document.add_paragraph(
    "The dataset contains customer purchasing information. "
    "It includes numerical features that help identify customer "
    "behaviour and purchasing patterns for segmentation."
)

# ====================================================
# Data Cleaning
# ====================================================

document.add_heading("4. Data Cleaning", level=1)

document.add_paragraph(
    "The dataset was inspected for missing values, duplicate records "
    "and incorrect data types. Missing values were handled and duplicate "
    "records were removed before further analysis."
)

# ====================================================
# Exploratory Data Analysis (EDA)
# ====================================================

document.add_heading("5. Exploratory Data Analysis", level=1)

document.add_paragraph(
    "Exploratory Data Analysis (EDA) was performed to understand the "
    "distribution of data and relationships among variables. "
    "Histograms, Boxplots and Correlation Heatmaps were generated."
)

# ====================================================
# Feature Scaling
# ====================================================

document.add_heading("6. Feature Scaling", level=1)

document.add_paragraph(
    "Feature Scaling was performed using StandardScaler to ensure that "
    "all numerical features contributed equally during clustering."
)

# ====================================================
# Principal Component Analysis
# ====================================================

document.add_heading("7. Principal Component Analysis (PCA)", level=1)

document.add_paragraph(
    "Principal Component Analysis (PCA) was used to reduce the "
    "dimensionality of the dataset while preserving most of the "
    "important information."
)

# ====================================================
# Elbow Method
# ====================================================

document.add_heading("8. Elbow Method", level=1)

document.add_paragraph(
    "The Elbow Method was used to determine the optimal number of "
    "clusters by analysing inertia values."
)

# ====================================================
# Silhouette Score
# ====================================================

document.add_heading("9. Silhouette Score", level=1)

document.add_paragraph(
    "Silhouette Score was calculated to evaluate the quality of "
    "the clusters. Higher scores indicate better clustering."
)

# ====================================================
# K-Means Clustering
# ====================================================

document.add_heading("10. K-Means Clustering", level=1)

document.add_paragraph(
    "K-Means Clustering grouped customers into similar clusters "
    "based on their purchasing behaviour."
)

# ====================================================
# Business Insights
# ====================================================

document.add_heading("11. Business Insights", level=1)

document.add_paragraph(
    "- High-value customers should receive loyalty rewards.\n"
    "- Medium-value customers can be targeted with personalised offers.\n"
    "- Low-value customers can be re-engaged through marketing campaigns.\n"
    "- Customer segmentation helps improve marketing efficiency."
)
    # ====================================================
# Add Graphs
# ====================================================

graph_list = [
    ("Histogram", "outputs/03_eda/histograms.png"),
    ("Correlation Heatmap", "outputs/03_eda/correlation_heatmap.png"),
    ("Explained Variance", "outputs/05_pca/explained_variance.png"),
    ("Cumulative Variance", "outputs/05_pca/cumulative_variance.png"),
    ("PCA Scatter Plot", "outputs/05_pca/pca_scatter_plot.png"),
    ("Elbow Method", "outputs/06_clustering/elbow_method.png"),
    ("Customer Clusters", "outputs/06_clustering/customer_clusters.png"),
    ("Cluster Distribution", "outputs/06_clustering/cluster_distribution.png"),
    ("Silhouette Score", "outputs/07_evaluation/silhouette_score_plot.png")
]

document.add_heading("13. Project Graphs", level=1)

for title, image in graph_list:
    if os.path.exists(image):
        document.add_heading(title, level=2)
        document.add_picture(image, width=Inches(5.8))

# ====================================================
# Save Report
os.makedirs("../outputs/09_report", exist_ok=True)

document.save("outputs/09_report/DecodeLabs_Project3_Report.docx")

print("Professional report generated successfully!")